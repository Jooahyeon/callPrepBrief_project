# -*- coding: utf-8 -*-
"""
콜 프렙 브리프(.md) → 디자인이 적용된 Word(.docx) / PDF 변환기.

사용법:
    python scripts/render_brief.py docs/briefs/브리프_삼성SDS_20260616.md
    # → output/브리프_삼성SDS_20260616.docx, .pdf 생성

고정 포맷 본문을 구조(제목/메타/섹션/Q&A/목록)로 파싱해 글꼴·색상·여백으로
한눈에 분류되도록 서식을 입힌다. 내용은 원본 .md 와 100% 동일하게 유지한다.
"""
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from fpdf import FPDF

# ---- 디자인 팔레트 (RGB) ----
NAVY = (31, 56, 100)     # 제목/섹션 헤더 글자
HEADER_BG = (31, 56, 100)  # 섹션 헤더 배경
META_BG = (221, 230, 241)  # 메타정보 박스 배경
Q_COLOR = (192, 80, 77)    # 질문 강조(붉은 계열)
A_COLOR = (64, 64, 64)     # 답변 글자
TEXT = (33, 33, 33)
MUTED = (110, 110, 110)

FONT = r"C:\Windows\Fonts\malgun.ttf"
FONT_BD = r"C:\Windows\Fonts\malgunbd.ttf"
FONT_NAME = "맑은 고딕"


# ============================================================
# 1) 파싱: 고정 포맷 본문 → (title, meta[], sections[(name, lines[])])
# ============================================================
def _is_sep(s):
    return len(s) >= 3 and set(s) <= set("=")


def _is_header(s):
    return s.startswith("[") and s.endswith("]")


def parse_brief(text):
    lines = text.split("\n")
    title, meta, sections = "", [], []
    cur = None
    state = "title"
    for raw in lines:
        s = raw.strip()
        if _is_sep(s):
            continue
        if state == "title":
            if not s:
                continue
            title = s.strip("[] ").strip()
            state = "meta"
            continue
        if _is_header(s):
            cur = [s.strip("[] ").strip(), []]
            sections.append(cur)
            state = "section"
            continue
        if state == "meta":
            if s:
                meta.append(s)
        elif state == "section":
            cur[1].append(raw.rstrip())
    # 섹션 앞뒤 빈 줄 제거
    for sec in sections:
        body = sec[1]
        while body and not body[0].strip():
            body.pop(0)
        while body and not body[-1].strip():
            body.pop()
    return title, meta, sections


def split_kv(line):
    """'고객사: 삼성SDS' → ('고객사', '삼성SDS')"""
    m = re.match(r"^\s*([^:：]+)\s*[:：]\s*(.*)$", line)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return None, line.strip()


def split_label(s):
    """'1. 현황 파악 — 설명' → ('1. 현황 파악', ' — 설명').

    번호 목록에서 em-dash(—/–/-) 앞의 라벨을 볼드 강조하기 위한 분리.
    구분자가 없으면 (전체, '') 를 돌려준다.
    """
    for sep in (" — ", " – ", " - "):
        if sep in s:
            lead, rest = s.split(sep, 1)
            return lead, sep + rest
    return s, ""


# ============================================================
# 2) Word 렌더링
# ============================================================
def _shade(paragraph, rgb):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "%02X%02X%02X" % rgb)
    paragraph._p.get_or_add_pPr().append(shd)


def _set_run(run, *, size=10.5, bold=False, color=TEXT):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def render_docx(title, meta, sections, path):
    doc = Document()
    sec = doc.sections[0]
    # 2페이지 이내로 담기 위한 좁은 여백
    sec.left_margin = sec.right_margin = Pt(40)
    sec.top_margin = sec.bottom_margin = Pt(34)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p.add_run(title), size=15, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(5)

    # 메타 정보 박스 (표)
    if meta:
        table = doc.add_table(rows=0, cols=2)
        table.allow_autofit = True
        for line in meta:
            k, v = split_kv(line)
            row = table.add_row().cells
            kp = row[0].paragraphs[0]
            kp.paragraph_format.space_after = Pt(0)
            _set_run(kp.add_run(k or ""), size=9, bold=True, color=NAVY)
            _shade(kp, META_BG)
            row[0].width = Pt(85)
            vp = row[1].paragraphs[0]
            vp.paragraph_format.space_after = Pt(0)
            _set_run(vp.add_run(v), size=9, color=TEXT)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def labeled(p, s, size, color):
        """'1. 라벨 — 설명' 에서 라벨만 볼드 처리."""
        lead, rest = split_label(s)
        _set_run(p.add_run(lead), size=size, bold=True, color=color)
        if rest:
            _set_run(p.add_run(rest), size=size, bold=False, color=color)

    # 섹션
    for name, body in sections:
        # 섹션 헤더 바
        hp = doc.add_paragraph()
        _shade(hp, HEADER_BG)
        hp.paragraph_format.space_before = Pt(6)
        hp.paragraph_format.space_after = Pt(2)
        r = hp.add_run("  " + name)
        _set_run(r, size=11, bold=True, color=(255, 255, 255))

        for line in body:
            s = line.strip()
            if not s:
                continue
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(1)
            pf.line_spacing = 1.0
            if re.match(r"^Q\d+\.", s):          # 질문
                _set_run(p.add_run(s), size=9.5, bold=True, color=Q_COLOR)
                pf.space_before = Pt(3)
            elif re.match(r"^A\d+\.", s):        # 답변
                _set_run(p.add_run(s), size=9, color=A_COLOR)
                pf.left_indent = Pt(12)
            elif re.match(r"^\d+\.", s):         # 번호 목록 (라벨 볼드)
                labeled(p, s, 9.5, TEXT)
                pf.left_indent = Pt(12)
            elif s.startswith("-"):              # 불릿 / 링크
                _set_run(p.add_run("• " + s.lstrip("- ").strip()),
                         size=8.5, color=TEXT)
                pf.left_indent = Pt(12)
            else:
                _set_run(p.add_run(s), size=9.5, color=TEXT)

    doc.save(path)


# ============================================================
# 3) PDF 렌더링
# ============================================================
class BriefPDF(FPDF):
    def header(self):
        pass


def _md_escape(s):
    """fpdf2 markdown 모드에서 의도치 않은 강조를 막기 위해 * 이스케이프."""
    return s.replace("*", r"\*")


def render_pdf(title, meta, sections, path):
    pdf = BriefPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.set_margins(12, 12, 12)
    pdf.add_page()
    pdf.add_font("malgun", "", FONT)
    pdf.add_font("malgun", "B", FONT_BD)
    W = pdf.w - pdf.l_margin - pdf.r_margin

    # 제목
    pdf.set_font("malgun", "B", 15)
    pdf.set_text_color(*NAVY)
    pdf.cell(0, 9, title, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # 메타 박스
    if meta:
        pdf.set_fill_color(*META_BG)
        key_w = 28
        for line in meta:
            k, v = split_kv(line)
            y0 = pdf.get_y()
            pdf.set_font("malgun", "", 9)
            pdf.set_text_color(*TEXT)
            pdf.set_xy(pdf.l_margin + key_w, y0)
            pdf.multi_cell(W - key_w, 6, " " + v, new_x="LMARGIN", new_y="NEXT")
            y1 = pdf.get_y()
            pdf.set_font("malgun", "B", 9)
            pdf.set_text_color(*NAVY)
            pdf.set_xy(pdf.l_margin, y0)
            pdf.cell(key_w, y1 - y0, " " + (k or ""), fill=True, border=0)
            pdf.set_xy(pdf.l_margin, y1)
        pdf.ln(2)

    for name, body in sections:
        # 섹션 헤더 바
        if pdf.get_y() > pdf.h - 24:
            pdf.add_page()
        pdf.set_fill_color(*HEADER_BG)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("malgun", "B", 11)
        pdf.ln(1)
        pdf.cell(0, 7, "  " + name, fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

        for line in body:
            s = line.strip()
            if not s:
                continue
            if re.match(r"^Q\d+\.", s):
                pdf.set_font("malgun", "B", 9.5)
                pdf.set_text_color(*Q_COLOR)
                pdf.multi_cell(W, 5, s, new_x="LMARGIN", new_y="NEXT")
            elif re.match(r"^A\d+\.", s):
                pdf.set_font("malgun", "", 9)
                pdf.set_text_color(*A_COLOR)
                pdf.set_x(pdf.l_margin + 4)
                pdf.multi_cell(W - 4, 4.6, s, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(0.6)
            elif re.match(r"^\d+\.", s):           # 번호 목록 (라벨 볼드)
                lead, rest = split_label(s)
                pdf.set_font("malgun", "", 9.5)
                pdf.set_text_color(*TEXT)
                pdf.set_x(pdf.l_margin + 3)
                txt = "**" + _md_escape(lead) + "**" + _md_escape(rest)
                pdf.multi_cell(W - 3, 5, txt, new_x="LMARGIN", new_y="NEXT",
                               markdown=True)
            elif s.startswith("-"):
                pdf.set_font("malgun", "", 8.5)
                pdf.set_text_color(*TEXT)
                pdf.set_x(pdf.l_margin + 3)
                pdf.multi_cell(W - 3, 4.8, "• " + s.lstrip("- ").strip(),
                               new_x="LMARGIN", new_y="NEXT")
            else:
                pdf.set_font("malgun", "", 9.5)
                pdf.set_text_color(*TEXT)
                pdf.multi_cell(W, 5, s, new_x="LMARGIN", new_y="NEXT")

    pdf.output(path)
    return len(pdf.pages)


def main(md_path):
    with open(md_path, encoding="utf-8") as f:
        text = f.read()
    title, meta, sections = parse_brief(text)
    base = os.path.splitext(os.path.basename(md_path))[0]
    os.makedirs("output", exist_ok=True)
    docx_path = os.path.join("output", base + ".docx")
    pdf_path = os.path.join("output", base + ".pdf")
    render_docx(title, meta, sections, docx_path)
    render_pdf(title, meta, sections, pdf_path)
    print("DOCX:", os.path.abspath(docx_path))
    print("PDF :", os.path.abspath(pdf_path))


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("usage: python scripts/render_brief.py <brief.md>")
    main(sys.argv[1])
